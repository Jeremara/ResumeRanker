import os
import streamlit as st
import pandas as pd
from io import BytesIO
from PyPDF2 import PdfReader
from dotenv import load_dotenv
from langchain.chat_models import AzureChatOpenAI
from langchain.prompts import PromptTemplate
from docx import Document
from fpdf import FPDF

# Load environment variables
load_dotenv()

# Configure Azure OpenAI
llm = AzureChatOpenAI(
    openai_api_base=os.getenv("AZURE_OPENAI_API_BASE"),
    openai_api_version=os.getenv("AZURE_OPENAI_API_VERSION"),
    openai_api_key=os.getenv("AZURE_OPENAI_API_KEY"),
    model_name=os.getenv("AZURE_OPENAI_MODEL_NAME"),
    temperature=0.9,
)

# Initialize session state
if "downloaded" not in st.session_state:
    st.session_state.downloaded = False

st.image("logo.jpg", width=100) 
st.title("Miah AI Resume Ranker")
st.markdown("Upload a **Job Description** and up to **10 Resumes (PDFs)** to rank them using AzureOpenAI.")

# --- Upload JD ---
jd_file = st.file_uploader("Upload Job Description (PDF or Text)", type=["pdf", "txt"])
jd_text = ""

if jd_file:
    if jd_file.type == "application/pdf":
        reader = PdfReader(jd_file)
        jd_text = "\n".join([page.extract_text() for page in reader.pages])
    else:
        jd_text = jd_file.read().decode("utf-8")

# --- Upload Resumes ---
resume_files = st.file_uploader("Upload up to 10 and more Resumes (PDF)", type="pdf", accept_multiple_files=True)
resume_texts = {}

results = []

if jd_text and resume_files:
    st.info("Processing resumes and ranking...")

    # Extract text from resumes
    for file in resume_files:
        reader = PdfReader(file)
        resume_text = "\n".join([page.extract_text() for page in reader.pages])
        resume_texts[file.name] = resume_text

    with st.spinner("Ranking resumes..."):
        for filename, text in resume_texts.items():
            prompt = PromptTemplate(
                input_variables=["job_description", "resume_text"],
                template="""
You are a hiring assistant. Compare the resume against the job description.

Job Description:
{job_description}

Resume:
{resume_text}

Rate this resume out of 100 based on how well it fits the job description. Be strict and unbiased.

Respond ONLY with a number (e.g., 78).
"""
            )

            formatted_prompt = prompt.format(
                job_description=jd_text[:3000],
                resume_text=text[:3000]
            )

            try:
                score_response = llm.predict(formatted_prompt)
                score = int(''.join(filter(str.isdigit, score_response)))
                results.append((filename, score))
            except Exception as e:
                results.append((filename, 0))

    # Sort by score
    results.sort(key=lambda x: x[1], reverse=True)

    # Display Results
    st.subheader("🏆 Ranked Resumes")
    for name, score in results:
        st.markdown(f"**{name}**: {score}/100")

    df_results = pd.DataFrame(results, columns=["Resume", "Score"])

    # --- User selects download format ---
    st.markdown("###  Download Ranked Results")
    format_option = st.selectbox("Choose file format", ["CSV", "DOCX"])

    buffer = BytesIO()
    filename = f"ranked_resumes.{format_option.lower()}"

    if format_option == "CSV":
        buffer.write(df_results.to_csv(index=False).encode())
    elif format_option == "DOCX":
        doc = Document()
        doc.add_heading("Resume Rankings", level=1)
        for name, score in results:
            doc.add_paragraph(f"{name}: {score}/100")
        doc.save(buffer)

     # Download + Reset trigger
    if not st.session_state.downloaded:
        if st.download_button(
            label=f"Download as {format_option}",
            data=buffer.getvalue(),
            file_name=filename,
            mime="application/octet-stream",
            key="download_button"
        ):
            st.session_state.downloaded = True
            #st.rerun()
    elif  st.session_state.downloaded:
        # Clear session to reset app
        st.session_state.clear()
        #st.rerun()
        
    # Reset Button
    if st.button("🔄 Reset and Upload New Document"):
        st.session_state.clear()
        st.rerun()